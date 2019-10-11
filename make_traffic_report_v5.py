# coding:utf-8
'''
@author: Sam_Shen
@file: make_traffic_report.py
@time: 2019/10/11 15:22
@desc:

需求:自动化生成流量报告

自定义:
1) 开始时间datestart
2) 结束时间dateend
3) 9.30上的url

实现效果:
1) 爬取SVR上的流量图片
2) 图像识别流量数值分析top5(部分识别有误的需人工干预)
3) 流量图和top5导入word模板后生成docx

2019.10.10 更新功能:
1) 图像打开,手动点击获取截图像素点坐标
2) 更新机器识别流量值的判定机制

2019.10.11 更新功能:
1) 进行了tesserocr的训练
2) 修正没有日期对应url的bug
'''


import datetime
import os
import re
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from PIL import Image
import tesserocr
import psutil
import operator
import cv2
import numpy as np

class Report:
    def __init__(self):
        self.base_url = "http://172.20.9.30/20191011/graph_10346.html"
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        self.img_dir = os.path.join(self.base_dir, 'img')
        self.template_dir = os.path.join(self.base_dir,'template')

        # 以下需根据要求手动调整
        self.Paging = 3                 # 正文显示: 几个流量图分页
        self.date_start = "20190901"    # 开始时间
        self.date_end = "20190930"      # 结束时间


    # 一些初始化工作
    def preparation(self):
        # 没有img目录就创建目录,有img目录则清空目录下的文件
        if not os.path.isdir(self.img_dir):
            os.mkdir(self.img_dir)
        else:
            for root, dirs, files in os.walk(self.img_dir, topdown=False):
                if files:
                    for name in files:
                        os.remove(os.path.join(root, name))

        if not os.path.isfile(os.path.join(self.template_dir,"model.docx")):
            print('未在本地发现template目录或model.docx模板文件,请检查.')
            exit(1)


    # 转换url
    # http://172.20.9.30/20190720/graph_12377.html 转换成 http://172.20.9.30/12377/graphs/graph_12377_1.png
    def trans_url(self):
        base_url_list = re.split('[/_.]', self.base_url)
        pic_id = base_url_list[-2]
        pic_date = base_url_list[-4]
        pic_url = "http://172.20.9.30/%s/graphs/graph_%s_1.png" % (pic_date,pic_id)
        return pic_url


    # 生成日期列表
    def create_assist_date(self):
        # 转为日期格式
        datestart = datetime.datetime.strptime(self.date_start,'%Y%m%d')
        dateend = datetime.datetime.strptime(self.date_end,'%Y%m%d')
        date_list = []
        date_list.append(datestart.strftime('%Y%m%d'))
        while datestart < dateend:
            # 日期叠加一天
            datestart += datetime.timedelta(days=+1)
            # 日期转字符串存入列表
            date_list.append(datestart.strftime('%Y%m%d'))
        return date_list


    # 生成图片url列表
    def create_pic_url(self):
        date_key = pic_url.split('/')[3]
        pic_list = []
        for date_str in date_list:
            ''' pic_dict = {'date': '20190901', 'url': 'http://172.20.9.30/20190902/graphs/graph_12328_1.png'} '''
            pic_dict = {}
            date_obj = datetime.datetime.strptime(date_str, '%Y%m%d')
            date_obj += datetime.timedelta(days=+1)
            date_add = date_obj.strftime("%Y%m%d")
            new_url = pic_url.replace(date_key, date_add)
            pic_dict["date"] = date_str
            pic_dict["url"] = new_url
            pic_list.append(pic_dict)
        return pic_list


    # 下载图片
    def download_pic(self):
        for each in pic_list:
            res = requests.get(each['url'])
            if res.status_code == 200:
                img = res.content
                img_name = os.path.join(self.img_dir,each['date']+'.png')
                with open(img_name,'wb' ) as f:
                    f.write(img)


    # 获取需要图片截图所需的像素点坐标
    def get_pixel_coordinate(self):
        img_file = os.path.join(self.img_dir,self.date_start+'.png')
        img = cv2.imdecode(np.fromfile(img_file, dtype=np.uint8), cv2.IMREAD_UNCHANGED)
        a = []
        b = []

        def on_EVENT_LBUTTONDOWN(event, x, y, flags, param):
            if event == cv2.EVENT_LBUTTONDOWN:
                xy = "%d,%d" % (x, y)
                a.append(x)
                b.append(y)
                cv2.circle(img, (x, y), 1, (255, 0, 0), thickness=-1)
                cv2.putText(img, xy, (x, y), cv2.FONT_HERSHEY_PLAIN, 1.0, (0, 0, 0), thickness=1)
                cv2.imshow("image", img)

        h, w, l = img.shape
        cv2.namedWindow("image", cv2.WINDOW_NORMAL)
        cv2.resizeWindow("image", w * 2, h * 2)
        try:
            cv2.setMouseCallback("image", on_EVENT_LBUTTONDOWN)
            cv2.imshow("image", img)
            cv2.waitKey(0)
            box = (a[0], b[0], a[1], b[1])
            return box
        except IndexError as e:
            print(e,"未在图片上选择任何坐标点")
            exit(1)


    # 识别图片中的最大流量,部分识别有误的需人工干预
    def realize_traffic(self):
        traffic_list = []

        # 关闭打开图片的process
        def close_photo():
            for proc in psutil.process_iter():
                if proc.name() == "Microsoft.Photos.exe":
                    proc.kill()

        # 判断机器识别流量精准度,返回False为识别错误需要人工干预
        def judge_traffic(val):
            val = val.replace(" ", "")

            # 如果首位不为数字
            if val[0].isdigit() == False:
                # print('首位不为数字')
                return False

            # 如果末位没有单位
            if val[-1] not in ["G", "M", "k"]:
                # print('末位没有单位或单位错误')
                return False

            # 如果末2位为.
            if val[-2] in ["."]:
                # print('末2位为.')
                return False

            # 如果每个切片单位为数字或为.
            sum = 0
            for i in val[:-1]:
                if i.isdigit():
                    pass
                elif i in ["."]:
                    pass
                else:
                    # print('切片单位不为数字或.')
                    sum += 1

            if sum == 0:
                return True
            elif sum >= 1:
                return False


        for root, dirs, files in os.walk(self.img_dir, topdown=False):
            for name in files:
                traffic_dict = {}
                img_files = os.path.join(root, name)
                img = Image.open(img_files)

                # 设置图像裁剪区域 (x左上，y左上，x右下,y右下)
                # box = (530, 228, 610, 250)

                # 图像裁剪
                image = img.crop(box)
                image = image.convert("L")
                w, h = image.size
                image = image.resize((w * 2, h * 2))

                # 图像二值化
                threshold = 180
                table = []
                for i in range(256):
                    if i < threshold:
                        table.append(0)
                    else:
                        table.append(1)
                image = image.point(table, '1')

                traffic_dict['filename'] = name
                traffic = tesserocr.image_to_text(image,lang="num").strip('\n')
                print(traffic)

                # 人工干预识别图像
                if not judge_traffic(traffic):
                    image.show()
                    traffic = input('File %s machine realize is %s,input normal traffic: ' % (name, traffic))
                    close_photo()

                traffic_dict['traffic'] = traffic.replace(" ", "").upper()
                traffic_list.append(traffic_dict)
        return traffic_list


    # 根据图片分析top5
    def calc_top5(self):
        traffic_list_update = []

        for each in traffic_list:
            traffic_dict = {}
            traffic_dict['filename'] = each["filename"]

            traffic = each["traffic"]
            if traffic[-1:] in ["G", "g"]:
                traffic_val = float(traffic[:-1]) * 1000 * 1000
            if traffic[-1:] in ["M", "m"]:
                traffic_val = float(traffic[:-1]) * 1000
            if traffic[-1:] in ["K", "k"]:
                traffic_val = float(traffic[:-1])
            traffic_dict['traffic_val'] = traffic_val
            traffic_dict['traffic_show'] = traffic
            traffic_list_update.append(traffic_dict)

        sorted_x = sorted(traffic_list_update, key=operator.itemgetter('traffic_val'), reverse=True)
        top5 = sorted_x[0:5]

        '''
        [{'filename': '20190906.png', 'traffic_val': 203540.0, 'traffic_show': '203.54M'}, 
            {'filename': '20190907.png', 'traffic_val': 186640.0, 'traffic_show': '186.64M'}, 
            {'filename': '20190908.png', 'traffic_val': 177610.0, 'traffic_show': '177.61M'}, 
            {'filename': '20190909.png', 'traffic_val': 172100.0, 'traffic_show': '172.10M'},
            {'filename': '20190912.png', 'traffic_val': 166520.0, 'traffic_show': '166.52M'}
        ]
        '''
        return top5


    # 生成word
    def create_docx(self):
        document = Document(os.path.join(self.template_dir,"model.docx"))

        def datestr_format(date_str):
            date_obj = datetime.datetime.strptime(date_str, '%Y%m%d')
            date_str_format = date_obj.strftime("%Y/%m/%d")
            return date_str_format

        # 添加Top5标题
        date_start_str = datestr_format(self.date_start)
        date_end_str = datestr_format(self.date_end)
        top5_title = "%s - %s Top5" % (date_start_str, date_end_str)
        document.add_heading(top5_title, 0)

        # 添加top5表格信息
        table = document.add_table(rows=1, cols=3)
        table.style = 'Light List Accent 5'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '日期'
        hdr_cells[1].text = '流量图'
        hdr_cells[2].text = '最大值'

        for each in top5:
            row_cells = table.add_row().cells
            row_cells[0].text = str(each['filename'].strip('.png'))

            # 单元格里添加图片
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(os.path.join(self.img_dir,each['filename']), width=Inches(3))

            row_cells[2].text = str(each['traffic_show'])

        document.add_page_break()  # 添加换页符

        # 添加正文标题
        daily_title = "%s - %s Traffic Daily" % (date_start_str, date_end_str)
        document.add_heading(daily_title, 0)

        # 设置正文字体格式
        document.styles['Normal'].font.name = u'Arial Unicode MS'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial Unicode MS')
        document.styles['Normal'].font.size = Pt(12)

        # docx文件添加时间内容和图片
        def insert_content(date_title, img_files):
            # 添加时间
            p = document.add_paragraph(date_title)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加图片
            # 检查图片文件是否存在,存在添加不存在则写入文件不存在
            if os.path.exists(img_files):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run("")
                run.add_picture(img_files, width=Inches(5.5))
            else:
                p = document.add_paragraph('图片404错误未找到')
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加换行符
            # p = document.add_paragraph('\n')

        num = 0
        for each in date_list:
            name = each+'.png'
            img_files = os.path.join(self.img_dir,name)

            for each in traffic_list:
                if name == each['filename']:
                    each_traffic = each['traffic']

            date_title = datestr_format(name.split('.')[0]) + " - " + each_traffic

            num += 1

            # 设定每*个图表为一页
            if num % self.Paging != 0:
                insert_content(date_title, img_files)

            else:
                insert_content(date_title, img_files)
                document.add_page_break()  # 添加换页符

        # 保存文件
        new_report = "model_" + datetime.datetime.now().strftime('%Y%m%d_%H%M%S') + ".docx"
        document.save(new_report)


'''生成2019年9月数据'''
report = Report()

report.preparation()
pic_url = report.trans_url()
date_list = report.create_assist_date()
pic_list = report.create_pic_url()
report.download_pic()
box = report.get_pixel_coordinate()
traffic_list = report.realize_traffic()
top5 = report.calc_top5()
report.create_docx()